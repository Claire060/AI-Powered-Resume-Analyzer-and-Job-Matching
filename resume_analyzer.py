import os
import re
import io
from docx import Document
import zipfile
from typing import Dict, Any, List
import numpy as np
from collections import Counter
from flask import Flask, request, jsonify, send_from_directory
import uuid
from sentence_transformers import SentenceTransformer, util
from datetime import datetime
import torch
from sklearn.metrics.pairwise import cosine_similarity

# Try importing cosine similarity safely
try:
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    cosine_similarity = None


class EnhancedResumeAnalyzer:
    def __init__(self, model_path: str = "sentence-transformers/all-MiniLM-L6-v2"):
        """
        Initialize the resume analyzer with Sentence-BERT model
        """
        self.model = None
        try:
            from sentence_transformers import SentenceTransformer
            print(f"Loading Sentence-BERT model: {model_path}")
            self.model = SentenceTransformer(model_path)
            print("Sentence-BERT model loaded successfully")
        except ImportError:
            print("Sentence Transformers not installed. Run: pip install sentence-transformers")
        except Exception as e:
            print(f"Could not load model: {e}")

        # Common resume sections
        self.sections = {
            'contact': ['contact', 'address', 'phone', 'email', 'linkedin'],
            'summary': ['summary', 'objective', 'profile'],
            'experience': ['experience', 'work history', 'employment', 'work experience'],
            'education': ['education', 'academic', 'qualifications', 'degrees'],
            'skills': ['skills', 'technical skills', 'competencies', 'technologies'],
            'projects': ['projects', 'personal projects', 'portfolio'],
            'certifications': ['certifications', 'certificates', 'licenses']
        }

        # Skill categories
        self.skill_categories = {
            'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust'],
            'web': ['html', 'css', 'react', 'angular', 'vue', 'django', 'flask', 'node.js', 'express'],
            'database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle'],
            'cloud': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform'],
            'data_science': ['pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn', 'r', 'matplotlib'],
            'tools': ['git', 'jenkins', 'jira', 'confluence', 'linux', 'bash']
        }

    def extract_zip(self, zip_path: str, extract_folder: str) -> List[str]:
            """Extract all files from a ZIP archive"""
            extracted_files = []
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_folder)
                    extracted_files = [os.path.join(extract_folder, f) for f in zip_ref.namelist()]
                    print(f"Extracted {len(extracted_files)} files from ZIP archive.")
            except Exception as e:
                print(f"Error extracting ZIP file: {e}")
            return extracted_files

    def analyze_resume(self, file_path: str, job_description: str = None) -> Dict[str, Any]:
        """Main function to analyze a resume"""
        print(f"\nAnalyzing resume: {file_path}")

        if not os.path.exists(file_path):
            return {'error': f"File not found: {file_path}"}

        # Enhanced text extraction with detailed logging
        text = self.extract_text_from_file(file_path)
        print(f"Extraction result: {len(text) if text else 0} characters")

        # FIXED: Better text validation
        if not text:
            return {'error': 'No text extracted from file'}
        elif isinstance(text, str) and (text.startswith("Error") or text.startswith("Unsupported")):
            return {'error': text}
        elif isinstance(text, str) and len(text.strip()) < 50:
            return {'error': 'Extracted text is too short. The file might be scanned or corrupted.'}

        print("Text extracted successfully")

        # Basic analysis
        basic_analysis = self._basic_text_analysis(text)
        sections = self._detect_sections(text)
        skills = self._extract_skills(text)
        experience = self._analyze_experience(text)

        # Prepare the initial analysis dictionary
        analysis = {
            'status': 'success',
            'file_name': os.path.basename(file_path),
            'basic_analysis': basic_analysis,
            'sections_found': sections,
            'skills': skills,
            'experience': experience,
            'raw_text_preview': text[:500] + "..." if len(text) > 500 else text,
            'text_length': len(text)
        }

        # Job description matching (optional)
        job_matches = []
        if job_description and self.model:
            print("Performing job description matching...")
            matching_results = self._match_job_description(text, job_description)
            
            job_matches.append({
                'job_title': "Senior Developer",  # This should be dynamic based on your matching logic
                'company': "TechCorp",  # This should be dynamic based on your matching logic
                'score': matching_results['match_percentage'],
                'industry': "Technology",  # This should be dynamic based on your matching logic
                'location': "San Francisco",  # This should be dynamic based on your matching logic
                'salary': "$100k-$120k",  # This should be dynamic based on your matching logic
                'skills': matching_results['common_keywords']
            })
            
            # Add more job matches to this list as required
            # Example: Iterate over a list of possible job postings and match scores

            # Add job matches to the analysis dictionary
            analysis['job_matches'] = job_matches  # Add the job matches to the analysis dictionary

        # Return the final analysis result
        return analysis


    # -------------------------
    # SIMPLIFIED TEXT EXTRACTION
    # -------------------------

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF, DOCX, or TXT files"""
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f" Processing {file_ext.upper()} file: {os.path.basename(file_path)}")
        
        try:
            if file_ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                return self.extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                return f"Unsupported file format: {file_ext}. Please use PDF, DOCX, or TXT."
        except Exception as e:
            return f"Error reading file: {str(e)}"

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf (since we know it works)"""
        try:
            from pypdf import PdfReader
            text = ""
            with open(file_path, 'rb') as file:  
                reader = PdfReader(file)
                print(f" Number of pages: {len(reader.pages)}")
                    
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        print(f"   Page {i+1}: {len(page_text)} characters")
            
            if text.strip():
                print(f" PDF extraction successful: {len(text)} total characters")
                return text.strip()
            else:
                return "Error: No text could be extracted from PDF"
                
        except ImportError:
            return "PDF processing requires 'pypdf' package. Run: pip install pypdf"
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
        return text.strip()
    
    def analyze_zip(self, zip_path: str, extract_folder: str, job_description: str = None):
        """Analyze all resumes within a ZIP file"""
        extracted_files = self.extract_zip(zip_path, extract_folder)
        results = []
        
        # Process each extracted resume file
        for file_path in extracted_files:
            result = self.analyze_resume(file_path, job_description)
            results.append(result)
        
        return results

    # -------------------------
    # TEXT ANALYSIS METHODS
    # -------------------------

    def _basic_text_analysis(self, text: str) -> Dict[str, Any]:
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        non_empty_sentences = [s for s in sentences if s.strip()]
        return {
            'word_count': len(words),
            'character_count': len(text),
            'sentence_count': len(non_empty_sentences),
            'avg_sentence_length': len(words) / max(len(non_empty_sentences), 1),
            'unique_words': len(set(words)),
            'readability_score': self._calculate_readability(text)
        }

    def _calculate_readability(self, text: str) -> float:
        sentences = [s for s in re.split(r'[.!?]+', text) if s.strip()]
        words = text.split()
        if not sentences or not words:
            return 0
        avg_sentence_length = len(words) / len(sentences)
        avg_syllables_per_word = self._estimate_syllables_per_word(text)
        readability = 206.835 - (1.015 * avg_sentence_length) - (84.6 * avg_syllables_per_word)
        return round(max(0, min(100, readability)), 2)

    def _estimate_syllables_per_word(self, text: str) -> float:
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        if not words:
            return 0
        vowels = 'aeiouy'
        total_syllables = 0
        for word in words:
            count = 0
            prev_vowel = False
            for char in word:
                if char in vowels and not prev_vowel:
                    count += 1
                    prev_vowel = True
                else:
                    prev_vowel = False
            if word.endswith('e'):
                count = max(1, count - 1)
            total_syllables += max(1, count)
        return total_syllables / len(words)

    def _detect_sections(self, text: str) -> Dict[str, List[str]]:
        lines = text.split('\n')
        sections = {key: [] for key in self.sections.keys()}
        current_section = 'unknown'
        for line in lines:
            line = line.strip()
            if not line:
                continue
            line_lower = line.lower()
            section_found = False
            for section_name, keywords in self.sections.items():
                if any(keyword in line_lower for keyword in keywords) and len(line.split()) <= 5:
                    current_section = section_name
                    section_found = True
                    break
            if not section_found and current_section != 'unknown':
                sections[current_section].append(line)
        for section_name in sections:
            sections[section_name] = [line for line in sections[section_name] if line.strip()]
        return sections

    def _extract_skills(self, text: str) -> Dict[str, Any]:
        text_lower = text.lower()
        found_skills = {}
        for category, skills in self.skill_categories.items():
            matched = [skill for skill in skills if re.search(r'\b' + re.escape(skill) + r'\b', text_lower)]
            if matched:
                found_skills[category] = matched
        all_skills = [skill for skills in found_skills.values() for skill in skills]
        return {
            'by_category': found_skills,
            'total_count': len(all_skills),
            'all_skills': all_skills
        }

    def _analyze_experience(self, text: str) -> Dict[str, Any]:
        date_patterns = [
            r'\b(19|20)\d{2}\b',
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}',
            r'\b\d{1,2}/\d{4}\b'
        ]
        experience_indicators = [
            'years', 'experience', 'worked', 'employed', 'position', 'role',
            'senior', 'junior', 'lead', 'manager', 'director'
        ]
        text_lower = text.lower()
        sentences = re.split(r'[.!?]+', text)
        experience_text = ". ".join(
            s.strip() for s in sentences if any(ind in s.lower() for ind in experience_indicators)
        )
        return {
            'experience_sentences': experience_text,
            'date_mentions': len(re.findall('|'.join(date_patterns), text)),
            'experience_indicators': len([i for i in experience_indicators if i in text_lower])
        }

    def _match_job_description(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        if not self.model:
            return {'error': 'Model not loaded for matching'}

        if not resume_text.strip() or not job_description.strip():
            return {'error': 'Empty resume or job description'}

        try:
            resume_emb = self.model.encode([resume_text])[0]
            job_emb = self.model.encode([job_description])[0]
            if cosine_similarity:
                similarity = cosine_similarity([resume_emb], [job_emb])[0][0]
            else:
                similarity = float(np.dot(resume_emb, job_emb) /
                                   (np.linalg.norm(resume_emb) * np.linalg.norm(job_emb)))
            job_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', job_description.lower()))
            resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower()))
            common = job_words.intersection(resume_words)
            return {
                'similarity_score': round(float(similarity), 4),
                'match_percentage': round(float(similarity * 100), 2),
                'common_keywords': list(common),
                'keyword_match_count': len(common),
                'total_job_keywords': len(job_words)
            }
        except Exception as e:
            return {'error': f'Matching failed: {str(e)}'}
        
    def check_resume_structure(self, resume_text):
        """Check resume structure completeness"""
        sections = {
            'Contact Information': False,
            'Summary or Objective': False,
            'Skills': False,
            'Experience': False,
            'Education': False,
            'Certifications': False,
            'Projects': False
        }

        # Expanded keywords for section headers
        keywords = {
            'Contact Information': ['contact', 'email', 'phone', 'address', 'mobile', 'linkedin'],
            'Summary or Objective': ['summary', 'objective', 'overview', 'profile', 'about'],
            'Skills': ['skill', 'competency', 'proficiency', 'technical skills', 'abilities'],
            'Experience': ['experience', 'work history', 'professional experience', 'employment history', 'work experience'],
            'Education': ['education', 'academic background', 'qualifications', 'degree', 'university'],
            'Certifications': ['certification', 'certificate', 'certifications', 'license', 'qualified'],
            'Projects': ['project', 'projects', 'portfolio', 'work samples']
        }

        # Simple checks for section headers
        lines = resume_text.split('\n')
        for line in lines:
            for section, section_keywords in keywords.items():
                if any(keyword in line.lower() for keyword in section_keywords):
                    sections[section] = True

        # Calculate structure score
        structure_score = sum(sections.values()) / len(sections) * 100
        return structure_score, sections

    def calculate_section_similarity(self, text1, text2):
        """Calculate similarity between two text sections using Sentence-BERT"""
        if not self.model:
            return 0
            
        if not text1.strip() or not text2.strip():
            return 0
        
        vec1 = self.model.encode(text1, convert_to_tensor=True)
        vec2 = self.model.encode(text2, convert_to_tensor=True)
        
        similarity = util.pytorch_cos_sim(vec1, vec2).item()
        return similarity * 100

    def extract_sections_from_resume(self, resume_text):
        """Extract different sections from resume text"""
        sections = {
            'skills': '',
            'experience': '',
            'education': '',
            'summary': ''
        }
        
        lines = resume_text.split('\n')
        current_section = None
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Detect section headers
            if any(keyword in line_lower for keyword in ['skill', 'technical']):
                current_section = 'skills'
            elif any(keyword in line_lower for keyword in ['experience', 'work', 'employment']):
                current_section = 'experience'
            elif any(keyword in line_lower for keyword in ['education', 'academic', 'degree']):
                current_section = 'education'
            elif any(keyword in line_lower for keyword in ['summary', 'objective', 'profile']):
                current_section = 'summary'
            elif current_section and line.strip():
                # Add content to current section
                sections[current_section] += line + ' '
        
        return sections

    def calculate_experience(self, resume_text):
        """Calculate experience level and years from resume text"""
        # Simple pattern matching for years of experience
        years_pattern = r'(\d+)\+?\s*(?:years?|yrs?)'
        matches = re.findall(years_pattern, resume_text.lower())
        
        years = 0
        if matches:
            # Take the highest number found
            years = max([int(match) for match in matches if match.isdigit()])
        
        # Determine experience level
        if years >= 8:
            level = "Senior/Expert"
        elif years >= 5:
            level = "Mid-Senior"
        elif years >= 3:
            level = "Mid-Level"
        elif years >= 1:
            level = "Junior"
        else:
            level = "Entry-Level"
            
        return level, years

    def calculate_comprehensive_score(self, resume_text, job_description=""):
        """Calculate comprehensive resume score using multiple factors"""
        
        # Extract skills
        skills = self._extract_skills(resume_text)
        all_skills = skills.get('all_skills', [])
        
        # Calculate experience
        experience_level, years = self.calculate_experience(resume_text)
        
        # Calculate structure score
        structure_score, sections_present = self.check_resume_structure(resume_text)
        
        # Base scores
        base_score = min(70 + (years * 4) + (len(all_skills) * 3), 95)
        
        # If job description provided, calculate match scores
        if job_description:
            job_skills = self._extract_skills(job_description)
            job_all_skills = job_skills.get('all_skills', [])
            
            # Skill matching score
            matched_skills = set(all_skills) & set(job_all_skills)
            skill_match_score = len(matched_skills) / len(job_all_skills) * 100 if job_all_skills else 0
            
            # Semantic similarity scores
            skills_similarity = self.calculate_section_similarity(
                ' '.join(all_skills), 
                ' '.join(job_all_skills)
            )
            
            # Overall semantic similarity
            overall_similarity = self.calculate_section_similarity(resume_text, job_description)
            
            # Combined match score
            match_score = (
                skill_match_score * 0.4 +
                skills_similarity * 0.3 +
                overall_similarity * 0.3
            )
            
            missing_skills = list(set(job_all_skills) - set(all_skills))
        else:
            # Use default job matching
            match_result = self.find_best_job_match(resume_text)
            match_score = match_result.get("match_score", 50)
            missing_skills = match_result.get("missing_skills", [])
            skill_match_score = (len(set(all_skills) - set(missing_skills)) / len(set(all_skills) | set(missing_skills)) * 100) if all_skills else 0
        
        # Content quality score
        content_score = self.calculate_content_quality(resume_text, skills, years)
        
        # Final weighted score
        final_score = (
            base_score * 0.3 +
            match_score * 0.3 +
            content_score * 0.2 +
            structure_score * 0.1 +
            skill_match_score * 0.1
        )
        
        return {
            "score": round(final_score),
            "content_score": round(content_score),
            "structure_score": round(structure_score),
            "match_score": round(match_score),
            "skill_match_score": round(skill_match_score),
            "total_skills": len(all_skills),
            "experience_level": f"{years} years ({experience_level})",
            "skills_detected": all_skills,
            "missing_skills": missing_skills,
            "sections_present": sections_present,
            "analysis_details": {
                "experience_years": years,
                "skill_categories": len(skills.get('by_category', {})),
                "assessment": self.get_assessment(final_score)
            }
        }

    def analyze_resume_text(self, resume_text, job_description=""):
        """Main analysis function using comprehensive scoring"""
        # Get comprehensive score
        analysis_result = self.calculate_comprehensive_score(resume_text, job_description)
        
        # Generate suggestions
        suggestions = self.generate_suggestions(
            analysis_result["skills_detected"], 
            analysis_result["missing_skills"], 
            analysis_result["analysis_details"]["experience_years"],
            analysis_result["sections_present"]
        )
        
        analysis_result["suggestions"] = suggestions
        
        # Get best job match
        match_result = self.find_best_job_match(resume_text)
        analysis_result["best_match_job"] = match_result.get("job_title", "Software Developer")
        analysis_result["matched_skills"] = match_result.get("matched_skills", [])
        
        return analysis_result

    def calculate_content_quality(self, resume_text, skills, years):
        """Calculate content quality score based on various factors"""
        score = 50  # Base score
        
        # Length factor
        word_count = len(resume_text.split())
        if 300 <= word_count <= 800:
            score += 15
        elif word_count > 800:
            score += 10
        else:
            score -= 10
        
        # Skills factor
        total_skills = sum(len(skill_list) for skill_list in skills.get('by_category', {}).values())
        if total_skills >= 8:
            score += 20
        elif total_skills >= 5:
            score += 15
        elif total_skills >= 3:
            score += 10
        
        # Experience factor
        if years >= 3:
            score += 15
        elif years >= 1:
            score += 10
        
        # Structure factor - check for common resume sections
        structure_score, _ = self.check_resume_structure(resume_text)
        score += (structure_score / 100) * 10
        
        return min(score, 100)

    def generate_suggestions(self, skills, missing_skills, years, sections_present):
        """Generate personalized improvement suggestions"""
        suggestions = []
        
        # Skill-based suggestions
        if missing_skills:
            suggestions.append(f"Consider learning or highlighting these missing skills: {', '.join(missing_skills[:3])}")
        
        if not any(keyword in str(skills) for keyword in ["cloud", "aws", "docker"]):
            suggestions.append("Add cloud technology skills like AWS or Docker to improve your marketability")
        
        if len([s for s in skills if s in ["communication", "teamwork", "problem solving"]]) < 2:
            suggestions.append("Include more soft skills like communication, teamwork, and problem-solving")
        
        # Structure-based suggestions
        missing_sections = [section for section, present in sections_present.items() if not present]
        if missing_sections:
            suggestions.append(f"Add missing sections: {', '.join(missing_sections[:2])}")
        
        # Experience-based suggestions
        if years < 2:
            suggestions.append("Highlight personal projects and contributions to demonstrate practical skills")
        elif years > 4:
            suggestions.append("Emphasize leadership experience and mentoring capabilities")
        
        # General suggestions
        suggestions.extend([
            "Use quantifiable achievements (e.g., 'Improved performance by 20%') instead of responsibilities",
            "Include industry-specific keywords to improve searchability",
            "Add relevant certifications and training programs"
        ])
        
        return suggestions[:5]  # Return top 5 suggestions

    def get_assessment(self, score):
        """Get textual assessment based on score"""
        if score >= 85:
            return "Excellent - Strong candidate for most roles"
        elif score >= 75:
            return "Very Good - Competitive profile with good skills"
        elif score >= 65:
            return "Good - Solid foundation with some areas for improvement"
        elif score >= 55:
            return "Fair - Focus on developing key skills and experience"
        else:
            return "Needs Improvement - Significant development needed in multiple areas"

    def find_best_job_match(self, resume_text):
        """Find the best matching job role based on skills"""
        # Simple job role matching based on skills
        job_roles = {
            "Software Developer": ["python", "java", "javascript", "c++", "git"],
            "Data Scientist": ["python", "pandas", "numpy", "machine learning", "statistics"],
            "Web Developer": ["javascript", "html", "css", "react", "node.js"],
            "DevOps Engineer": ["aws", "docker", "kubernetes", "linux", "terraform"],
            "Data Analyst": ["sql", "excel", "python", "tableau", "statistics"]
        }
        
        skills = self._extract_skills(resume_text)
        all_skills = skills.get('all_skills', [])
        
        best_match = ""
        best_score = 0
        matched_skills = []
        
        for job, required_skills in job_roles.items():
            matched = set(all_skills) & set(required_skills)
            score = len(matched) / len(required_skills) * 100
            
            if score > best_score:
                best_score = score
                best_match = job
                matched_skills = list(matched)
        
        return {
            "job_title": best_match,
            "match_score": round(best_score),
            "matched_skills": matched_skills,
            "missing_skills": list(set(job_roles.get(best_match, [])) - set(matched_skills))
        }


# ----------------------------------------------------------
# Public wrapper for Flask integration
# ----------------------------------------------------------
def analyze_resume(file_path: str, job_description: str = None) -> dict:
    """
    Public interface for Flask or external code.
    Initializes the EnhancedResumeAnalyzer and runs analysis.
    """
    analyzer = EnhancedResumeAnalyzer()
    return analyzer.analyze_resume(file_path, job_description)